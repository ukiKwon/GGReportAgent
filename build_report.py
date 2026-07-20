# -*- coding: utf-8 -*-
"""서울시 5개 구청 사업제안 기획안 docx 생성 스크립트"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC_PATH = r"C:\claude_workspace\서울시_5개구청_사업제안기획안.docx"

FONT_NAME = "맑은 고딕"

def set_korean_font(run, size=11, bold=False, color=None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)

def add_heading(doc, text, level=1, size=16, color=(0x1F, 0x4E, 0x79)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_korean_font(run, size=size, bold=True, color=color)
    return p

def add_para(doc, text, size=10.5, bold=False, color=None, space_after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_korean_font(run, size=size, bold=bold, color=color)
    return p

def add_bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_korean_font(run, size=size)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        set_korean_font(run, size=9.5, bold=True, color=(0xFF, 0xFF, 0xFF))
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '2E5395')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shd)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            run = p.add_run(str(val))
            set_korean_font(run, size=9)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("─" * 60)
    set_korean_font(run, size=8, color=(0xAA, 0xAA, 0xAA))

# ============================================================
# 데이터 정의: 5개 구청
# ============================================================

districts = [
    {
        "name": "도봉구",
        "color": (0x1F, 0x4E, 0x79),
        "trust_score": 74,
        "budget_total": "약 20~30억원",
        "budget_ratio": "구 예산(7,919억원) 대비 약 0.25~0.4%",
        "key_problems": [
            "4개년계획 121개 사업 중 27%(약24개)가 홈페이지 검색으로 확인 불가",
            "민원 채널 4종 이상 파편화, '구청장에게 바란다'는 답변완료 표시 없이 처리경과 추적 곤란",
            "2026년 최다 민원: 창동역 민자역사 출입구·보행로 신설 요청 집중",
            "간판개선·언덕길 열선 등 소규모 생활사업은 예산 배정에도 검색 자체가 안 됨",
        ],
        "top_projects": [
            ("IT-2", "민원 처리 현황 통합 대시보드 & 알림톡 연동", "300~450백만원", "최우선"),
            ("IT-1", "도봉구 통합 사업정보 검색 플랫폼 구축", "450~600백만원", "최우선"),
            ("IT-3", "창동역 일대 공사구간 스마트 보행안전 시스템", "180~250백만원", "상"),
            ("FN-2", "창동역 공사기간 인근 상가 영업피해 완충 지원금", "300~500백만원", "상"),
        ],
        "verdict": "spec 조사에서 실제로 확인된 문제(검색 미확인 27%, 민원 채널 파편화, 창동역 민원 집중)를 충실히 근거로 사용했으며 인용 수치 대부분이 원문과 일치. 다만 IT-1이 기존 계획사업(5-3 스마트데이터 허브센터, 5,000백만원)과의 중복 검토 없이 완전 신규처럼 제안된 점이 최대 리스크로, 실행 전 사실확인 필요.",
    },
    {
        "name": "노원구",
        "color": (0x2E, 0x86, 0x4B),
        "trust_score": 88,
        "budget_total": "약 15~23억원",
        "budget_ratio": "구 예산(1조 3,625억원) 대비 약 0.11~0.17%",
        "key_problems": [
            "23개 표본 중 확인됨 65%이나, 다수가 구청 홈페이지가 아닌 의회 회의록(council.nowon.kr) 의존",
            "서브도메인 9개 이상 파편화(회의록/환경/공단/뉴스레터/법무/부동산/데이터광장 등)",
            "'구청장에게 바란다'는 처리 지연·비공개 마스킹 다수, 반면 구민고충처리위원회는 신속 처리",
            "재건축진단 비용지원(여론 1순위 35.9%)과 새활용센터(120억 투입)가 확인안됨",
        ],
        "top_projects": [
            ("IT-1", "노원구 통합 사업추진현황 검색 포털 구축", "400~550백만원", "최우선"),
            ("IT-2", "민원 채널 통합 대시보드 & 알림톡 연동", "300~450백만원", "최우선"),
            ("FN-1", "재건축진단 비용지원 홍보·신청 프로세스 정비 지원금", "300~450백만원", "상"),
            ("FN-2", "석계역 일대 소상공인 간판개선·보행환경 연계 지원금", "250~400백만원", "상"),
        ],
        "verdict": "IT 4건은 spec 사업목록에 없는 완전 신규, FN 3건은 기존 계획사업(재건축진단지원·석계역 간판개선·새활용센터)의 확장으로 명확히 구분되어 '기존 사업을 신규처럼 포장'한 사례가 발견되지 않음. 5개 구청 중 근거 인용 정확성이 가장 높은 수준.",
    },
    {
        "name": "광진구",
        "color": (0xB0, 0x2A, 0x37),
        "trust_score": 88,
        "budget_total": "약 15~24억원",
        "budget_ratio": "구 예산(8,536.71억원) 대비 약 0.18~0.28%",
        "key_problems": [
            "44개 표본 중 부분확인 11개 다수가 2026년 신규 개관 사업(청년아지트·천천히편의점 등)에 집중",
            "민원 게시판 7종 이상 분산, 채널마다 처리상태 표기 체계 자체가 상이",
            "2026년 최다 민원: 모아타운 등 정비사업 민원(구의2동 재개발 반대 등) + 교통·보행 제안",
            "모아타운 등 26개 정비사업이 '예산 미표기' 민간사업 구조라 공식 진행상황이 회의록에만 존재",
        ],
        "top_projects": [
            ("IT-1", "신설사업 전용페이지 패스트트랙 & 통합사업정보 검색", "350~500백만원", "최우선"),
            ("IT-2", "민원채널 처리상태 표준화 & 통합 조회 시스템", "300~450백만원", "최우선"),
            ("IT-3", "모아타운 등 정비사업 통합 진행상황 알림 플랫폼", "200~300백만원", "상"),
            ("FN-1", "모아타운 등 정비사업 예정지 주민 생활불편 완충 지원금", "300~500백만원", "상"),
        ],
        "verdict": "7개 사업 전부 spec 원문 대조 결과 수치·명칭·날짜 인용 오류 없음. 기존 사업(모아타운, 골목소통, 청년아지트, 천천히편의점)을 대체하지 않고 홍보·추적성·안정화를 보완하는 방식으로 설계되어 이미 실행 중인 사업의 신규 포장 사례 없음.",
    },
    {
        "name": "동대문구",
        "color": (0x6A, 0x3D, 0x9A),
        "trust_score": 82,
        "budget_total": "약 19~28억원",
        "budget_ratio": "구 예산(9,824.36억원) 대비 약 0.19~0.29%",
        "key_problems": [
            "2024~2026년 주요업무계획 실물 PDF 부재, 홍보 동영상(mp4)만 게시되는 구조적 문서 공백",
            "122개 사업 중 확인율 52%(부분확인7%, 확인안됨41%), 생활밀착사업일수록 확인율 급감",
            "2026년 최다 민원: 신답극동아파트 리모델링조합 관련 안전진단·행정명령 촉구(상위 10건 중 8건)",
            "재정자립도 등 핵심 재정지표가 홈페이지에서 확인 불가",
        ],
        "top_projects": [
            ("IT-1", "주요업무계획 상시공개 및 마스터DB 구축", "350~500백만원", "최우선"),
            ("IT-3", "민원 채널 통합 대시보드 & 처리현황 알림 서비스", "300~450백만원", "최우선"),
            ("FN-1", "노후 소규모 공동주택 안전진단·보수 매칭 지원금", "400~600백만원", "상"),
            ("IT-4", "노후공동주택 안전이슈 조기경보 시스템", "150~220백만원", "상"),
        ],
        "verdict": "정책문서 실물 부재라는 근본 문제를 정확히 포착해 IT-1의 핵심 근거로 삼은 점이 강점. 다만 IT-2가 spec에 없는 2026년 분야별 예산 배분을 있는 것처럼 병기한 오류, FN-2가 '확인안됨'을 '실행중'으로 과잉 해석한 부분은 수정이 필요.",
    },
    {
        "name": "동작구",
        "color": (0xC2, 0x6B, 0x00),
        "trust_score": 84,
        "budget_total": "약 17~25억원",
        "budget_ratio": "구 예산(9,331.3억원) 대비 약 0.18~0.27%",
        "key_problems": [
            "22개 표본 중 확인안됨9%+부분확인32%, 2026년 신규 개관사업(JUMP1·동작청년센터 등)에 집중",
            "'구청장에게 바란다'·'옴부즈만 민원 신청' 게시판이 외부에서 완전 차단(WebFetch 2회 시도 모두 빈 페이지)",
            "동작구의회 회의록(assembly.dongjak.go.kr) 의존 구조 - 구 자체 통합 추적 페이지 부재",
            "세대별 39개 중점사업 대부분 예산액이 카드형 설명에 명시되지 않음",
        ],
        "top_projects": [
            ("IT-2", "민원채널 접근성 복구 & 처리현황 공개 대시보드", "300~450백만원", "최우선"),
            ("IT-1", "신규사업 개시 알림형 통합 사업정보 페이지 구축", "350~500백만원", "최우선"),
            ("FN-2", "노후 소규모 공동주택·빌라 안전개선 지원 확대", "400~600백만원", "상"),
            ("IT-3", "동작구의회 회의록 연계 사업추적 검색 위젯", "120~180백만원", "상"),
        ],
        "verdict": "재정자립도(28.87%)·예산 수치 등 인용 정확도가 높고, 민원채널 완전 차단이라는 5개 구청 중 가장 심각한 접근성 문제를 정확히 짚어냄. 다만 FN-1의 재원 성격(조기집행 vs 신규재원) 모호성과 일부 제안부서명의 실존 여부 미확정은 보완 필요.",
    },
]

# ============================================================
# 문서 생성 시작
# ============================================================

doc = Document()

# 기본 스타일 설정
style = doc.styles['Normal']
style.font.name = FONT_NAME
style.font.size = Pt(10.5)
rPr = style.element.get_or_add_rPr()
rFonts = OxmlElement('w:rFonts')
rFonts.set(qn('w:eastAsia'), FONT_NAME)
rPr.append(rFonts)

# ---- 표지 ----
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(80)
run = title_p.add_run("서울시 5개 구청\n사업제안 기획안")
set_korean_font(run, size=30, bold=True, color=(0x1F, 0x4E, 0x79))

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_before = Pt(20)
run = sub_p.add_run("도봉구 · 노원구 · 광진구 · 동대문구 · 동작구")
set_korean_font(run, size=15, color=(0x44, 0x44, 0x44))

sub2_p = doc.add_paragraph()
sub2_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2_p.paragraph_format.space_before = Pt(10)
run = sub2_p.add_run("IT/디지털 기획 및 금전적 지원 사업 제안 · 타당성 검증 결과 포함")
set_korean_font(run, size=11, color=(0x66, 0x66, 0x66))

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.paragraph_format.space_before = Pt(200)
run = date_p.add_run("작성일: 2026년 7월 16일")
set_korean_font(run, size=10, color=(0x88, 0x88, 0x88))

doc.add_page_break()

# ---- 1페이지 요약 (구청별 1장씩, 총 5장) ----
add_heading(doc, "구청별 사업 제안 요약", level=1, size=20)
add_para(doc, "각 구청의 spec 조사(4개년/주요업무계획 분석·홈페이지 검색확인·2026년 민원게시판 현황)에서 실제로 확인된 문제에 근거해 도출한 핵심 제안사업과 검증 결과를 1페이지로 요약합니다.", size=9.5, color=(0x66,0x66,0x66))
doc.add_page_break()

for i, d in enumerate(districts):
    add_heading(doc, f"{d['name']}  |  1페이지 요약", level=1, size=18, color=d["color"])

    # 신뢰도 점수 박스
    score_p = doc.add_paragraph()
    score_p.paragraph_format.space_after = Pt(10)
    run = score_p.add_run(f"제안서 신뢰도 점수:  {d['trust_score']} / 100")
    set_korean_font(run, size=13, bold=True, color=d["color"])

    add_para(doc, f"제안 총예산: {d['budget_total']}  ({d['budget_ratio']})", size=10, bold=True)

    add_heading(doc, "spec 조사에서 확인된 핵심 문제", level=2, size=12, color=(0x33,0x33,0x33))
    for prob in d["key_problems"]:
        add_bullet(doc, prob, size=9.8)

    add_heading(doc, "핵심 제안사업 (우선순위 상위 4건)", level=2, size=12, color=(0x33,0x33,0x33))
    add_table(
        doc,
        ["코드", "사업명", "추정예산", "우선순위"],
        [[p[0], p[1], p[2], p[3]] for p in d["top_projects"]],
        col_widths=[1.6, 8.5, 3.0, 2.0],
    )

    add_heading(doc, "검증 총평", level=2, size=12, color=(0x33,0x33,0x33))
    add_para(doc, d["verdict"], size=9.8)

    if i < len(districts) - 1:
        doc.add_page_break()

doc.add_page_break()

# ============================================================
# 구청별 상세 챕터
# ============================================================

def add_detail_file(doc, filepath, size=9.3):
    """spec/plan 텍스트 파일을 문단 단위로 그대로 추가"""
    with open(filepath, encoding="utf-8") as f:
        content = f.read()
    for line in content.split("\n"):
        line = line.rstrip()
        if not line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        is_header = line.startswith("=====") or line.startswith("[") and line.endswith("]")
        is_bracket_title = line.strip().startswith("[") and "]" in line
        run = p.add_run(line)
        if line.strip().startswith("====="):
            set_korean_font(run, size=7, color=(0xBB,0xBB,0xBB))
        elif is_bracket_title:
            set_korean_font(run, size=size+1, bold=True, color=(0x1F,0x4E,0x79))
        else:
            set_korean_font(run, size=size)

district_dirs = {
    "도봉구": "dobong",
    "노원구": "nowon",
    "광진구": "gwangjin",
    "동대문구": "dongdaemun",
    "동작구": "dongjak",
}

file_labels = [
    ("00_제안개요_및_배경.txt", "1. 제안 개요 및 배경"),
    ("01_제안사업_요약표.txt", "2. 제안사업 요약표"),
    ("02_IT디지털기획_사업제안.txt", "3. IT/디지털 기획 사업 제안"),
    ("03_금전적지원_사업제안.txt", "4. 금전적 지원 사업 제안"),
    ("04_실행로드맵_및_기대효과.txt", "5. 실행 로드맵 및 기대효과"),
    ("05_검증결과.txt", "6. 타당성 검증 결과"),
]

for di, d in enumerate(districts):
    dirname = district_dirs[d["name"]]
    add_heading(doc, f"[{di+1}] {d['name']} 상세 기획안", level=1, size=20, color=d["color"])
    add_divider(doc)

    for fname, label in file_labels:
        fpath = fr"C:\claude_workspace\{dirname}\plan\{fname}"
        add_heading(doc, label, level=2, size=13, color=d["color"])
        add_detail_file(doc, fpath)
        add_divider(doc)

    if di < len(districts) - 1:
        doc.add_page_break()

doc.save(DOC_PATH)
print(f"Saved: {DOC_PATH}")
